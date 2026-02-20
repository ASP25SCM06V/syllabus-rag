# app.py
# Flask backend for the Syllabus RAG Q&A system
#
# No LLM used for generation — answers are extracted directly from the
# most relevant chunks using embedding similarity. This means:
#   - Zero hallucination (every word comes from your document)
#   - No ~600MB model download
#   - Faster responses
#   - Works for any syllabus
#
# usage:
#   python app.py
#   open http://localhost:5000

import os
import sys
import re
import numpy as np
import faiss
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from sentence_transformers import SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter
from werkzeug.utils import secure_filename

app = Flask(__name__, static_folder=".")
CORS(app)

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"docx", "txt"}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

rag_state = {
    "chunks": [],
    "index": None,
    "filename": None,
    "ready": False,
}

print("\nLoading embedding model...")
embed_model = SentenceTransformer("all-MiniLM-L6-v2")
print("Embedding model ready!")
print("\n" + "="*50)
print("Server ready! Open http://localhost:5000")
print("="*50 + "\n")


# ----- document loading -----

def load_document(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        from docx import Document
        doc = Document(path)

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                seen = []
                for cell in cells:
                    if cell not in seen:
                        seen.append(cell)
                if seen:
                    table_texts.append(" | ".join(seen))

        all_text = "\n".join(paragraphs)
        if table_texts:
            all_text += "\n\n--- Tables ---\n" + "\n".join(table_texts)
        return all_text

    elif ext == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ----- index builder -----

def build_index(file_path):
    print(f"\nIndexing: {file_path}")
    raw_text = load_document(file_path)
    print(f"Loaded: {len(raw_text)} characters")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=50,
        length_function=len
    )
    chunks = splitter.split_text(raw_text)
    print(f"Split into {len(chunks)} chunks")

    print("Embedding...")
    chunk_embeddings = embed_model.encode(chunks, show_progress_bar=True)

    embedding_dim = chunk_embeddings.shape[1]
    index = faiss.IndexFlatL2(embedding_dim)
    index.add(np.array(chunk_embeddings).astype("float32"))
    print(f"FAISS ready: {index.ntotal} vectors")

    rag_state["chunks"] = chunks
    rag_state["index"] = index
    rag_state["ready"] = True
    return len(chunks)


# ----- answer builder -----
# 
# strategy: retrieve top-k chunks by embedding similarity
# then clean and join them — the answer IS the relevant chunks
# we do light cleanup to make it readable

def build_answer(question, retrieved_chunks, scores):
    # L2 distance — lower = more similar
    # if best score is too high, nothing relevant was found
    best_score = scores[0]
    if best_score > 40:
        return "I don't have that information in the uploaded syllabus."

    # use the top 1-3 most relevant chunks based on score
    answer_chunks = []
    for i, (chunk, score) in enumerate(zip(retrieved_chunks, scores)):
        if score > 30 and i > 0:
            break  # stop adding chunks if they get much less relevant
        # clean up the chunk a bit
        cleaned = chunk.strip()
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)  # collapse excessive newlines
        answer_chunks.append(cleaned)

    answer = "\n\n".join(answer_chunks)

    # cap length
    if len(answer) > 800:
        answer = answer[:800].rsplit(' ', 1)[0] + "..."

    return answer


# ----- RAG pipeline -----

def run_rag(question, k=4):
    if not rag_state["ready"]:
        return {"error": "No document indexed. Please upload a syllabus first."}

    chunks = rag_state["chunks"]
    index  = rag_state["index"]

    q_embedding = embed_model.encode([question]).astype("float32")
    distances, indices = index.search(q_embedding, k)

    retrieved_chunks  = [chunks[i] for i in indices[0]]
    retrieved_indices = [int(i) for i in indices[0]]
    retrieved_scores  = [float(distances[0][j]) for j in range(k)]

    answer = build_answer(question, retrieved_chunks, retrieved_scores)

    print(f"Q: {question}")
    print(f"Scores: {[round(s,1) for s in retrieved_scores]}")
    print(f"A: {answer[:150]}...")

    return {
        "answer": answer,
        "chunks": [
            {
                "index": retrieved_indices[i],
                "text":  retrieved_chunks[i],
                "score": retrieved_scores[i]
            }
            for i in range(len(retrieved_chunks))
        ],
        "total_chunks": len(chunks)
    }


# ----- routes -----

@app.route("/")
def serve_frontend():
    return send_from_directory(".", "big-data-rag.html")

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": "Only .docx and .txt files are supported"}), 400

    filename  = secure_filename(file.filename)
    save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(save_path)

    try:
        chunk_count = build_index(save_path)
        rag_state["filename"] = filename
        return jsonify({"success": True, "filename": filename, "chunks": chunk_count})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/ask", methods=["POST"])
def ask_endpoint():
    data     = request.get_json()
    question = data.get("question", "").strip()
    if not question:
        return jsonify({"error": "No question provided"}), 400
    return jsonify(run_rag(question))

@app.route("/health")
def health():
    return jsonify({
        "status":         "ok",
        "ready":          rag_state["ready"],
        "filename":       rag_state["filename"],
        "chunks_indexed": len(rag_state["chunks"])
    })


if __name__ == "__main__":
    if len(sys.argv) > 1:
        try:
            build_index(sys.argv[1])
            rag_state["filename"] = os.path.basename(sys.argv[1])
        except Exception as e:
            print(f"Warning: could not load file: {e}")
    app.run(debug=False, port=5000)
